"""
DOCX converter – MVP version.

Spec format (minimal):
{
  "paragraphs":[
      { "text":"Some text", "style":"Heading 1", "font": { "name":"Arial", "size_pt":12 } }
  ]
}
"""

from io import BytesIO
from typing import Dict, Any

from docx import Document
from docx.shared import Pt

from . import REGISTRY

async def render(spec: Dict[str, Any], watermark: bool) -> bytes:
    doc = Document()

    # --- basic paragraphs ---
    for block in spec.get("paragraphs", []):
        para = doc.add_paragraph()
        run = para.add_run(block.get("text", ""))
        if "style" in block:
            para.style = block["style"]
        if font := block.get("font"):
            run.font.name = font.get("name")
            if size := font.get("size_pt"):
                run.font.size = Pt(size)

    # --- serialize to bytes ---
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# self-register
REGISTRY["docx"] = render

